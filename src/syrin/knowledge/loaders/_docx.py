"""DOCX loader using python-docx (fallback when Docling not used)."""

from __future__ import annotations

import asyncio
from collections.abc import Iterable
from pathlib import Path
from typing import Protocol

from syrin.knowledge._document import Document, DocumentMetadata


class _DocxCellLike(Protocol):
    """Protocol for python-docx Cell (text)."""

    @property
    def text(self) -> str:
        """Cell text."""
        ...


class _DocxRowLike(Protocol):
    """Protocol for python-docx Row (cells)."""

    @property
    def cells(self) -> Iterable[_DocxCellLike]:
        """Row cells."""
        ...


class _DocxTableLike(Protocol):
    """Protocol for python-docx Table (rows)."""

    @property
    def rows(self) -> Iterable[_DocxRowLike]:
        """Table rows."""
        ...


def _check_docx() -> None:
    """Ensure python-docx is installed."""
    try:
        import docx  # noqa: F401
    except ImportError as err:
        raise ImportError(
            "python-docx is required for DOCXLoader. Install with: pip install syrin[docx]"
        ) from err


def _try_docling(path: Path) -> list[Document] | None:
    """Use DoclingLoader if docling is available. Returns None if not or on error."""
    try:
        from syrin.knowledge.loaders._docling import DoclingLoader

        loader = DoclingLoader(path)
        return loader.load()
    except ImportError:
        return None
    except Exception:
        return None


def _table_to_markdown(table: _DocxTableLike) -> str:
    """Convert python-docx Table to markdown string."""
    data: list[list[str]] = []
    for table_row in table.rows:
        cells = [cell.text.replace("|", "\\|").replace("\n", " ") for cell in table_row.cells]
        data.append(cells)

    if not data:
        return ""

    lines: list[str] = []
    for i, row_vals in enumerate(data):
        lines.append("| " + " | ".join(row_vals) + " |")
        if i == 0:
            lines.append("|" + "|".join("---" for _ in row_vals) + "|")
    return "\n".join(lines)


class DOCXLoader:
    """Load DOCX files using python-docx.

    Uses DoclingLoader when docling is installed (better table extraction).
    Otherwise uses python-docx for text and table extraction.

    Requires: pip install syrin[docx] (or syrin[docling] for best quality)
    """

    def __init__(self, path: str | Path, *, use_docling: bool = True) -> None:
        """Initialize DOCXLoader.

        Args:
            path: Path to the DOCX file.
            use_docling: If True, try Docling first when available.
        """
        self._path = Path(path)
        self._use_docling = use_docling

    @property
    def path(self) -> Path:
        """Path to the DOCX file."""
        return self._path

    def load(self) -> list[Document]:
        """Load the DOCX file. Uses Docling when available, else python-docx.

        Returns:
            List of Documents. Main body text first, then table Documents
            with table_markdown in metadata when applicable.
        """
        if not self._path.exists():
            raise FileNotFoundError(f"DOCX file does not exist: {self._path}")

        if self._use_docling:
            docs = _try_docling(self._path)
            if docs is not None:
                return docs

        _check_docx()
        from docx import Document as DocxDocument

        doc = DocxDocument(str(self._path))
        result: list[Document] = []

        text_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if text_parts:
            result.append(
                Document(
                    content="\n\n".join(text_parts),
                    source=str(self._path),
                    source_type="docx",
                    metadata={"section": "body"},
                )
            )

        for idx, table in enumerate(doc.tables):
            md = _table_to_markdown(table)
            meta: DocumentMetadata = {
                "table_index": idx,
                "table_markdown": md,
            }
            result.append(
                Document(
                    content=md,
                    source=str(self._path),
                    source_type="docx_table",
                    metadata=meta,
                )
            )

        if not result:
            result.append(
                Document(
                    content="",
                    source=str(self._path),
                    source_type="docx",
                    metadata={},
                )
            )

        return result

    async def aload(self) -> list[Document]:
        """Load DOCX asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load)


__all__ = ["DOCXLoader"]
