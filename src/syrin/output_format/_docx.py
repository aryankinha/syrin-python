"""DOCX output formatter via python-docx."""

from __future__ import annotations

from io import BytesIO
from typing import cast

_DOCX_EXTRA_HINT = "Install with: pip install syrin[docx]"


def _get_document_class() -> type:
    """Lazy import Document from python-docx. Raises ImportError with hint if not installed."""
    try:
        from docx import Document

        return cast(type, Document)
    except ImportError as e:
        raise ImportError(
            f"python-docx is required for DOCX output. {_DOCX_EXTRA_HINT}. Original: {e}"
        ) from e


class DOCXFormatter:
    """Formatter for Word (.docx) documents. Requires syrin[docx]."""

    def format(self, content: str, *, metadata: dict[str, object] | None = None) -> bytes:
        """Convert content to DOCX bytes."""
        Document = _get_document_class()
        doc = Document()

        title = ""
        if metadata and "title" in metadata:
            t = metadata["title"]
            if isinstance(t, str) and t:
                title = t
                doc.add_heading(title, 0)

        for para in content.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def extension(self) -> str:
        return "docx"
